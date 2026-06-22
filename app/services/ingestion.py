from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, BinaryIO

import pandas as pd
import yaml
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".csv", ".xlsx", ".xls", ".docx", ".json", ".yaml", ".yml"}


@dataclass(frozen=True)
class ParsedDocument:
    source_name: str
    content: str
    metadata: dict[str, Any]


class DocumentParsingError(RuntimeError):
    """Raised when an uploaded document cannot be parsed into text."""


def _read_bytes(file_obj: Any) -> bytes:
    try:
        if hasattr(file_obj, "getvalue"):
            data = file_obj.getvalue()
        elif hasattr(file_obj, "read"):
            current_position = file_obj.tell() if hasattr(file_obj, "tell") else None
            data = file_obj.read()
            if current_position is not None and hasattr(file_obj, "seek"):
                file_obj.seek(current_position)
        else:
            with Path(file_obj).open("rb") as handle:
                data = handle.read()
    except Exception as exc:
        raise DocumentParsingError(f"Unable to read uploaded file bytes: {exc}") from exc

    if not isinstance(data, bytes):
        raise DocumentParsingError("Uploaded file did not provide bytes")
    if not data:
        raise DocumentParsingError("Uploaded file is empty")
    return data


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParsingError("Unable to decode text using supported encodings")


def _clean_text(text: str) -> str:
    normalized_lines = [line.strip() for line in text.replace("\x00", " ").splitlines()]
    return "\n".join(line for line in normalized_lines if line).strip()


def _parse_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"Page {index}\n{page_text}")
        return _clean_text("\n\n".join(pages))
    except Exception as exc:
        raise DocumentParsingError(f"PDF parsing failed: {exc}") from exc


def _parse_txt(data: bytes) -> str:
    try:
        return _clean_text(_decode_text(data))
    except Exception as exc:
        raise DocumentParsingError(f"TXT parsing failed: {exc}") from exc


def _parse_csv(data: bytes) -> str:
    try:
        text = _decode_text(data)
        frame = pd.read_csv(StringIO(text))
        return _clean_text(frame.to_csv(index=False))
    except Exception as exc:
        raise DocumentParsingError(f"CSV parsing failed: {exc}") from exc


def _parse_xlsx(data: bytes) -> str:
    try:
        workbook = pd.read_excel(BytesIO(data), sheet_name=None)
        sections = []
        for sheet_name, frame in workbook.items():
            sections.append(f"Sheet: {sheet_name}\n{frame.to_csv(index=False)}")
        return _clean_text("\n\n".join(sections))
    except Exception as exc:
        raise DocumentParsingError(f"XLSX parsing failed: {exc}") from exc


def _parse_docx(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        return _clean_text("\n".join(paragraphs + table_rows))
    except Exception as exc:
        raise DocumentParsingError(f"DOCX parsing failed: {exc}") from exc


def _parse_json(data: bytes) -> str:
    try:
        loaded = json.loads(_decode_text(data))
        return _clean_text(json.dumps(loaded, ensure_ascii=True, indent=2, sort_keys=True))
    except Exception as exc:
        raise DocumentParsingError(f"JSON parsing failed: {exc}") from exc


def _parse_yaml(data: bytes) -> str:
    try:
        loaded = yaml.safe_load(_decode_text(data))
        return _clean_text(yaml.safe_dump(loaded, sort_keys=True, allow_unicode=False))
    except Exception as exc:
        raise DocumentParsingError(f"YAML parsing failed: {exc}") from exc


def parse_uploaded_file(file_obj: BinaryIO | Any) -> ParsedDocument:
    source_name = getattr(file_obj, "name", None) or str(file_obj)
    extension = Path(source_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParsingError(f"Unsupported file type '{extension}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}")

    data = _read_bytes(file_obj)
    parsers = {
        ".pdf": _parse_pdf,
        ".txt": _parse_txt,
        ".csv": _parse_csv,
        ".xlsx": _parse_xlsx,
        ".xls": _parse_xlsx,
        ".docx": _parse_docx,
        ".json": _parse_json,
        ".yaml": _parse_yaml,
        ".yml": _parse_yaml,
    }
    content = parsers[extension](data)
    if not content:
        raise DocumentParsingError(f"No readable text could be extracted from {source_name}")

    return ParsedDocument(
        source_name=Path(source_name).name,
        content=content,
        metadata={"source": Path(source_name).name, "extension": extension, "bytes": len(data)},
    )
